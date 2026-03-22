import os
import json
import re
from pathlib import Path
from typing import Dict, Any

# 可选导入
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    Document = None

# ========== 文件提取 ==========
def extract_text_from_file(file_path: str) -> str:
    """从多种格式文件中提取文本"""
    ext = Path(file_path).suffix.lower()
    if ext in ['.txt', '.md']:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif ext == '.pdf':
        if not PyPDF2:
            raise ImportError("请安装 PyPDF2: pip install PyPDF2")
        text = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return "\n".join(text)
    elif ext == '.docx':
        if not Document:
            raise ImportError("请安装 python-docx: pip install python-docx")
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        raise ValueError(f"不支持的文件格式: {ext}")

# ========== 进度管理 ==========
def load_progress(report_id: str, resume: bool) -> Dict[str, Any]:
    """加载进度文件，若 resume=False 则返回初始进度"""
    progress_file = f"progress_{report_id}.json"
    if resume and os.path.exists(progress_file):
        with open(progress_file, 'r', encoding='utf-8') as f:
            return json.load(f)
    else:
        return {
            "report_id": report_id,
            "completed_chapters": [],
            "output_file": f"report_{report_id}.md"
        }

def save_progress(progress: Dict[str, Any]):
    """保存进度到文件"""
    progress_file = f"progress_{progress['report_id']}.json"
    with open(progress_file, 'w', encoding='utf-8') as f:
        json.dump(progress, f, indent=2, ensure_ascii=False)

def append_to_markdown(output_file: str, section: Dict[str, Any], content: str):
    """将生成的内容追加到 Markdown 文件"""
    with open(output_file, 'a', encoding='utf-8') as f:
        f.write(f"## {section['id']} {section['title']}\n\n")
        f.write(content)
        f.write("\n\n")

# ========== Markdown 转 Word ==========
def unescape_markdown(text: str) -> str:
    """去除可能由 JSON 转义引入的反斜杠，例如 \* -> *"""
    text = text.replace('\\*', '*')
    return text

def set_chinese_font(doc):
    """设置文档默认字体为宋体"""
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    style.font.size = Pt(12)

    for i in range(1, 10):
        heading_style = doc.styles.get(f'Heading {i}')
        if heading_style:
            heading_style.font.name = '宋体'
            heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            heading_style.font.size = Pt(16 - i)

def clear_paragraph(paragraph):
    """清空段落的所有 run"""
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)

def parse_inline_formatting(paragraph, text):
    """解析 Markdown 行内格式（粗体）并添加到段落"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            if part:
                paragraph.add_run(part)

def _add_text_to_cell(cell, text, bold=False):
    """向表格单元格添加文本，支持行内格式"""
    cell.text = ''
    p = cell.add_paragraph()
    parse_inline_formatting(p, text)
    if bold:
        for run in p.runs:
            run.bold = True

def _parse_table(doc, table_lines):
    """解析 Markdown 表格并添加到文档"""
    table_lines = [ln.strip() for ln in table_lines if ln.strip()]
    if len(table_lines) < 3:
        return

    header = [cell.strip() for cell in table_lines[0].strip('|').split('|')]
    align_line = table_lines[1].strip('|')
    aligns = []
    for cell in align_line.split('|'):
        cell = cell.strip()
        if cell.startswith(':') and cell.endswith(':'):
            aligns.append('center')
        elif cell.endswith(':'):
            aligns.append('right')
        else:
            aligns.append('left')
    aligns = (aligns + ['left'] * len(header))[:len(header)]

    table = doc.add_table(rows=1 + len(table_lines)-2, cols=len(header))
    table.style = 'Light Grid Accent 1'

    # 表头
    hdr_cells = table.rows[0].cells
    for idx, cell_text in enumerate(header):
        _add_text_to_cell(hdr_cells[idx], cell_text, bold=True)

    # 数据行
    for row_idx, data_line in enumerate(table_lines[2:]):
        cells = [cell.strip() for cell in data_line.strip('|').split('|')]
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(cells):
            if col_idx < len(row_cells):
                _add_text_to_cell(row_cells[col_idx], cell_text)
                for paragraph in row_cells[col_idx].paragraphs:
                    paragraph.alignment = {
                        'left': WD_ALIGN_PARAGRAPH.LEFT,
                        'center': WD_ALIGN_PARAGRAPH.CENTER,
                        'right': WD_ALIGN_PARAGRAPH.RIGHT
                    }.get(aligns[col_idx], WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()

def parse_markdown_to_docx(doc: Document, text: str):
    """将 Markdown 文本解析并添加到 Word 文档"""
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # 处理标题
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            content = line.lstrip('#').strip()
            heading = doc.add_heading(level=min(level, 9))
            clear_paragraph(heading)
            parse_inline_formatting(heading, content)
            i += 1
            continue

        # 处理表格
        if line.startswith('|') and i + 1 < len(lines) and re.match(r'^[\|\-\:\s]+$', lines[i+1].strip()):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            _parse_table(doc, table_lines)
            continue

        # 处理无序列表
        if line and line[0] in ('-', '*', '+') and line[1] == ' ':
            p = doc.add_paragraph(style='List Bullet')
            content = line[2:].strip()
            parse_inline_formatting(p, content)
            i += 1
            continue

        # 处理普通段落
        if line.strip():
            p = doc.add_paragraph()
            parse_inline_formatting(p, line)
        else:
            doc.add_paragraph('')
        i += 1

def convert_markdown_to_docx(markdown_content: str, output_docx: str):
    """将 Markdown 字符串转换为 Word 文档并保存"""
    doc = Document()
    set_chinese_font(doc)
    parse_markdown_to_docx(doc, markdown_content)
    doc.save(output_docx)